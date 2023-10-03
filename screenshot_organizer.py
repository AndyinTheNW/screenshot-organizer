import os
from docx import Document
from docx.shared import Inches
from PIL import Image

parent_folder = r'E:\OneDrive - Bravo Project & Outsourcing\Documentos\ShareX\Screenshots\2023-10'

for subdirectory in os.listdir(parent_folder):

    subdirectory_path = os.path.join(parent_folder, subdirectory)


    if os.path.isdir(subdirectory_path):

        screenshot_files = [
            os.path.join(subdirectory_path, filename)
            for filename in sorted(os.listdir(subdirectory_path), key=lambda x: os.path.getmtime(os.path.join(subdirectory_path, x)))
        ]

        doc = Document()
      
        for screenshot_path in screenshot_files:

            if screenshot_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):

                doc.add_picture(screenshot_path, width=Inches(6)) 

                doc.add_paragraph()
            else:
                print(f"Skipping unsupported file: {screenshot_path}")

        output_docx = os.path.basename(subdirectory_path) + '.docx'
        doc.save(os.path.join(subdirectory_path, output_docx))
